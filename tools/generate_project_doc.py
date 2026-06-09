from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "项目说明.docx"
DOC_PATH = DOCS / "项目说明.doc"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 45, 58)
MUTED = RGBColor(86, 98, 112)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, bottom=80, start=120, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tbl_cell_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    set_cell_margins(table)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.167


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("星轨迷阵 Star Maze")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Java Swing 小游戏项目说明")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_numbered(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_table(doc: Document, headers, rows, widths) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], LIGHT_FILL)
        for paragraph in hdr[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_code_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(50, 65, 82)


def build_docx() -> None:
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("一、项目概述", level=1)
    doc.add_paragraph(
        "《星轨迷阵》是一款基于 Java Swing 开发的桌面迷宫冒险小游戏。玩家操控探索者在程序生成的霓虹迷宫中收集星核，"
        "躲避巡航体追踪，并在收集足够星核后启动传送门进入下一层。项目不依赖第三方游戏引擎，主要使用 Swing 自绘界面、"
        "Timer 游戏循环、键盘输入映射、音频合成和本地存档完成完整游戏体验。"
    )

    add_table(
        doc,
        ["项目项", "说明"],
        [
            ["项目名称", "星轨迷阵 Star Maze"],
            ["开发语言", "Java，使用 JDK 17 及以上可编译运行"],
            ["界面技术", "Java Swing / AWT 自绘界面"],
            ["发布方式", "可运行 Jar：dist/StarMaze.jar"],
            ["主类", "com.starmaze.StarMazeApp"],
        ],
        [1.5, 5.0],
    )

    doc.add_heading("二、功能清单", level=1)
    add_bullets(
        doc,
        [
            "完整开始界面、暂停界面、帮助界面、设置界面、过关界面和失败界面。",
            "随机迷宫关卡生成，层数提升后地图、敌人、星核、补给和裂隙都会变化。",
            "星核收集、传送门锁定/解锁、分数奖励、移动扣分、关卡奖励和最高分记录。",
            "巡航体敌人具有巡逻和近距离追踪行为，并会在相位干扰后短暂眩晕。",
            "相位穿墙机制：消耗能量短时间穿越墙体，也可用来干扰敌人。",
            "轨迹回溯机制：把玩家和敌人拉回几秒前的位置，提供逃生和策略空间。",
            "量子裂隙格提供随机折跃或能量补充，增加路线选择的不确定性。",
            "本地存档记录最高分、游玩次数、通关次数和音效开关。",
            "程序生成音效，无需额外资源文件，包含移动、收集、相位、回溯、胜利和失败提示。",
        ],
    )

    doc.add_heading("三、创新性与可玩性设计", level=1)
    add_table(
        doc,
        ["创新点", "玩法价值"],
        [
            ["相位穿墙", "玩家不只是沿迷宫路线移动，可以主动穿越墙体改变路径，提升解谜和逃生自由度。"],
            ["轨迹回溯", "在被追踪或路线失误时可回到几秒前，形成风险管理和二次决策。"],
            ["动态敌人 AI", "敌人平时巡逻，接近玩家后追踪，使每局节奏更紧张。"],
            ["随机裂隙", "裂隙可能折跃也可能补能，玩家需要判断是否冒险踩入。"],
            ["程序生成关卡", "每次新游戏和每层地图不同，减少重复感。"],
        ],
        [1.45, 5.05],
    )

    doc.add_heading("四、类结构说明", level=1)
    add_table(
        doc,
        ["类/模块", "主要职责"],
        [
            ["StarMazeApp", "程序入口，设置 Swing 外观并创建主窗口。"],
            ["GameFrame / GamePanel", "窗口和 Swing 调度入口，帧循环、渲染模块编排、计时器停止和测试用循环状态检查。"],
            ["InputController / InputActions / MenuMouseController / RenderClock", "键盘绑定、键盘菜单动作、菜单点击、鼠标悬停光标控制和渲染帧节奏管理。"],
            ["MenuMetrics / MenuPanel", "菜单面板尺寸、居中位置和按钮布局几何参数。"],
            ["GameViewRenderer / HudRenderer / GameSceneRenderer / MenuOverlayRenderer / MenuRenderer 等", "视图渲染门面以及 HUD、棋盘场景、实体、背景、特效和菜单的独立绘制模块，其中 MenuOverlayRenderer 负责菜单模式分发，HudRenderer 委托 HudPanelRenderer 绘制 HUD 面板底板、HudStatsRenderer 绘制标题与分数统计、HudMeterRenderer 绘制能量条、HudRewindRenderer 绘制回溯次数、HudMessageRenderer 绘制右侧提示消息，EntityRenderer 继续细分为 PickupRenderer、ExitRenderer、PlayerRenderer 和 EnemyRenderer，MenuRenderer 继续委托 MenuChromeRenderer 绘制菜单遮罩和面板底板、TitleMenuContentRenderer 绘制标题内容、SimpleMenuContentRenderer 绘制暂停/设置内容、MenuButtonGroupRenderer 组装按钮文案并委托 MenuButtonRenderer 绘制菜单按钮、HelpMenuContentRenderer 绘制帮助内容、ResultMenuContentRenderer 绘制结算内容，EffectRenderer 继续委托 RingEffectRenderer 绘制环形特效、StunEffectRenderer 绘制眩晕点特效，场景渲染器维护静态棋盘缓存。"],
            ["GameMessages / MenuText / HudText", "集中维护游戏内提示语、菜单标题、帮助说明、结果面板文案和 HUD 标签。"],
            ["GameState", "核心游戏状态、移动规则、得分、相位、回溯、过关和失败逻辑。"],
            ["ScoreRules / PhaseRules / RewindRules / LevelProgressRules / RiftRules / RiftOutcome", "拾取物、裂隙、相位、回溯、通关奖励、星级评级、相位能量上限/回复、回溯次数边界、量子裂隙结果以及关卡推进奖励等纯规则计算。"],
            ["EnemyController / PathFinder", "敌人巡逻、A* 追踪决策和路径搜索。"],
            ["LevelPopulator", "星核、相位电池、回溯补给和敌人出生点填充。"],
            ["RewindHistory / MenuButtonCache", "轨迹回溯快照管理、菜单按钮布局缓存和鼠标命中测试。"],
            ["Level / LevelGenerator", "地图数据结构与随机迷宫生成，包括通路、墙、裂隙和出口。"],
            ["Enemy / Position / Direction / TileType", "实体、坐标、方向和地图格类型等基础模型。"],
            ["SaveData", "最高分、统计数据和设置的本地持久化。"],
            ["SoundEngine / SoundEffect", "使用 Java Sound API 合成短音效。"],
            ["StarMazeSmokeTest", "无界面自检入口，验证关卡、敌人、星核、计时和基础技能逻辑。"],
        ],
        [2.05, 4.45],
    )

    doc.add_heading("五、核心算法与实现", level=1)
    doc.add_heading("1. 随机迷宫生成", level=2)
    doc.add_paragraph(
        "LevelGenerator 先以深度优先回溯算法雕刻连通迷宫，再按关卡等级增加额外回路，减少单一路线带来的枯燥感。"
        "出口通过广度优先搜索选取距离起点最远的可达格，使玩家需要完成较完整的探索过程。"
    )
    doc.add_heading("2. 敌人巡逻与追踪", level=2)
    doc.add_paragraph(
        "敌人在普通状态下沿当前方向巡逻，遇墙或阻挡时随机换向。当玩家进入警戒范围，EnemyController 会调用 PathFinder 的 A* "
        "启发式搜索寻找敌人到玩家的下一步方向。相位状态下敌人不会追踪玩家，碰撞时敌人会被眩晕，从而把技能设计与 AI 行为关联起来。"
    )
    doc.add_heading("3. 轨迹回溯", level=2)
    doc.add_paragraph(
        "游戏每隔数帧记录玩家位置、敌人位置和相位剩余时间。玩家按 R 后消耗回溯次数，恢复到历史快照，"
        "并让敌人短暂失衡。该机制既能纠错，又能制造主动诱敌后回撤的策略。"
    )

    doc.add_heading("六、操作说明", level=1)
    add_table(
        doc,
        ["按键", "功能"],
        [
            ["WASD / 方向键", "控制玩家上下左右移动。"],
            ["Space", "启动相位穿墙，消耗相位能量。"],
            ["R", "轨迹回溯。"],
            ["P / Esc", "暂停或返回。"],
            ["H", "打开玩法说明。"],
            ["M", "切换音效。"],
            ["Enter", "标题界面开始游戏，过关/失败界面继续。"],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("七、编译与发布", level=1)
    doc.add_paragraph("项目根目录提供 build.ps1，可完成编译和 Jar 打包。")
    add_code_line(doc, r".\build.ps1")
    add_code_line(doc, r"java -jar .\dist\StarMaze.jar")
    doc.add_paragraph("发布目录 dist 中包含 StarMaze.jar 和 run.bat，Windows 下可双击 run.bat 启动。")

    doc.add_heading("八、测试结果", level=1)
    add_bullets(
        doc,
        [
            "已执行 build.ps1，javac 编译成功并生成 dist/StarMaze.jar。",
            "已执行 jar 内无界面自检：java \"-Djava.awt.headless=true\" -cp dist\\StarMaze.jar com.starmaze.StarMazeSmokeTest。",
            "自检输出为 StarMaze smoke test passed，覆盖关卡生成、敌人生成、星核数量、计时更新、相位技能、基础移动状态、确定性收集星核并过关流程、确定性敌人碰撞失败/相位反制流程、确定性轨迹回溯流程、确定性相位结束安全恢复流程、EnemyController 追踪/巡逻/眩晕/无路回退、PathFinder A* 最短路/占用/不可达边界、Level 导航辅助方法、LevelGenerator 可达性/尺寸上限/裂隙位置、LevelPopulator 数量上限和安全落点、ScoreRules 分数和评级边界、SaveData 正常保存和损坏存档恢复、InputController 键盘绑定安装、BoardRenderer 静态缓存/外框/裂隙绘制、BackgroundRenderer 背景绘制、StarField 星点绘制与尺寸边界、PickupRenderer 拾取物绘制、ExitRenderer 锁定/开启出口绘制、PlayerRenderer 普通/相位玩家绘制、EnemyRenderer 方向/眩晕敌人绘制、MenuMouseController 鼠标悬停/点击事件、GameViewRenderer 渲染门面、MenuOverlayRenderer 菜单覆盖层分发、MenuChromeRenderer 遮罩和面板绘制、HudPanelRenderer 面板底板绘制、HudStatsRenderer 统计文字绘制、HudMeterRenderer 能量条绘制、HudRewindRenderer 回溯次数绘制、HudMessageRenderer 右侧提示绘制、FooterRenderer 底部提示绘制、TitleMenuContentRenderer 标题内容绘制、SimpleMenuContentRenderer 暂停/设置内容绘制、MenuButtonGroupRenderer 按钮组绘制、MenuButtonRenderer 按钮绘制、HelpMenuContentRenderer 帮助内容绘制、ResultMenuContentRenderer 结算内容绘制、RingEffectRenderer 环形特效绘制、StunEffectRenderer 眩晕点特效绘制、离屏 GamePanel 非空渲染和计时器启停检查。",
            "Jar 清单包含 Main-Class: com.starmaze.StarMazeApp，可通过 java -jar dist/StarMaze.jar 启动游戏窗口。",
        ],
    )

    doc.add_heading("九、提交文件", level=1)
    add_numbered(
        doc,
        [
            "源码：src/main/java/com/starmaze 下的 Java 源码及 build.ps1、README.md。",
            "项目说明：docs/项目说明.doc，同时保留 docs/项目说明.docx 便于新版 Word 编辑。",
            "项目发布 Jar：dist/StarMaze.jar，另附 dist/run.bat 便于双击运行。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("星轨迷阵 Star Maze - Java Swing 课程项目")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(DOCX_PATH)


def build_html_doc() -> None:
    sections = [
        ("一、项目概述", [
            "《星轨迷阵》是一款基于 Java Swing 开发的桌面迷宫冒险小游戏。玩家操控探索者在程序生成的霓虹迷宫中收集星核，躲避巡航体追踪，并在收集足够星核后启动传送门进入下一层。",
            "项目不依赖第三方游戏引擎，主要使用 Swing 自绘界面、Timer 游戏循环、键盘输入映射、音频合成和本地存档完成完整游戏体验。",
            "主类为 com.starmaze.StarMazeApp，发布文件为 dist/StarMaze.jar。",
        ]),
        ("二、功能清单", [
            "完整开始、暂停、帮助、设置、过关和失败界面。",
            "随机迷宫关卡生成，层数提升后地图、敌人、星核、补给和裂隙都会变化。",
            "星核收集、传送门锁定/解锁、分数奖励、最高分记录和本地存档。",
            "巡航体敌人具有巡逻和近距离追踪行为。",
            "相位穿墙、轨迹回溯、量子裂隙和程序生成音效提升可玩性。",
        ]),
        ("三、创新性与可玩性设计", [
            "相位穿墙：玩家可短时间穿越墙体，同时干扰巡航体，改变传统迷宫只能沿路移动的限制。",
            "轨迹回溯：记录玩家与敌人的历史位置，关键时刻恢复到数秒前，形成纠错和策略反制。",
            "动态敌人 AI：敌人平时巡逻，接近玩家后通过路径搜索追踪。",
            "量子裂隙：踩入后可能折跃或补充能量，使路线选择更有风险与收益。",
            "程序生成关卡：每局地图、道具和敌人位置不同，提升重复游玩价值。",
        ]),
        ("四、类结构说明", [
            "StarMazeApp：程序入口，创建主窗口。",
            "GameFrame / GamePanel：窗口和 Swing 调度入口，负责帧循环、渲染模块编排、计时器停止和测试用循环状态检查。",
            "InputController / InputActions / MenuMouseController / RenderClock：键盘绑定、键盘菜单动作、菜单点击、鼠标悬停光标控制和渲染帧节奏管理。",
            "MenuMetrics / MenuPanel：菜单面板尺寸、居中位置和按钮布局几何参数。",
            "GameViewRenderer / HudRenderer / GameSceneRenderer / MenuOverlayRenderer / MenuRenderer 等：视图渲染门面以及 HUD、棋盘场景、实体、背景、特效和菜单的独立绘制模块，其中 MenuOverlayRenderer 负责菜单模式分发，HudRenderer 委托 HudPanelRenderer 绘制 HUD 面板底板、HudStatsRenderer 绘制标题与分数统计、HudMeterRenderer 绘制能量条、HudRewindRenderer 绘制回溯次数、HudMessageRenderer 绘制右侧提示消息，EntityRenderer 继续细分为 PickupRenderer、ExitRenderer、PlayerRenderer 和 EnemyRenderer，MenuRenderer 继续委托 MenuChromeRenderer 绘制菜单遮罩和面板底板、TitleMenuContentRenderer 绘制标题内容、SimpleMenuContentRenderer 绘制暂停/设置内容、MenuButtonGroupRenderer 组装按钮文案并委托 MenuButtonRenderer 绘制菜单按钮、HelpMenuContentRenderer 绘制帮助内容、ResultMenuContentRenderer 绘制结算内容，EffectRenderer 继续委托 RingEffectRenderer 绘制环形特效、StunEffectRenderer 绘制眩晕点特效，场景渲染器维护静态棋盘缓存。",
            "GameMessages / MenuText / HudText：集中维护游戏内提示语、菜单标题、帮助说明、结果面板文案和 HUD 标签。",
            "GameState：核心游戏状态、移动规则、得分、相位、回溯、过关和失败逻辑。",
            "ScoreRules / PhaseRules / RewindRules / LevelProgressRules / RiftRules / RiftOutcome：拾取物、裂隙、相位、回溯、通关奖励、星级评级、相位能量上限/回复、回溯次数边界、量子裂隙结果以及关卡推进奖励等纯规则计算。",
            "EnemyController / PathFinder：敌人巡逻、A* 追踪决策和路径搜索。",
            "LevelPopulator：星核、相位电池、回溯补给和敌人出生点填充。",
            "RewindHistory / MenuButtonCache：轨迹回溯快照管理、菜单按钮布局缓存和鼠标命中测试。",
            "Level / LevelGenerator：地图结构与随机迷宫生成。",
            "Enemy / Position / Direction / TileType：实体、坐标、方向和地图格类型。",
            "SaveData：最高分、统计数据和设置持久化。",
            "SoundEngine / SoundEffect：程序合成短音效。",
            "StarMazeSmokeTest：无界面自检入口。",
        ]),
        ("五、核心算法与实现", [
            "随机迷宫生成：使用深度优先回溯算法雕刻连通迷宫，并按关卡等级加入额外回路。",
            "出口选择：通过广度优先搜索选取距离起点较远的可达格。",
            "敌人追踪：玩家进入警戒范围后，敌人使用 A* 启发式搜索计算朝向玩家的下一步。",
            "BFS 使用场景：出口选择和相位结束后的安全格恢复仍使用广度优先搜索。",
            "轨迹回溯：定期保存历史快照，按 R 后恢复玩家和敌人位置并让敌人短暂眩晕。",
        ]),
        ("六、操作说明", [
            "WASD 或方向键：移动。",
            "Space：启动相位穿墙。",
            "R：轨迹回溯。",
            "P/Esc：暂停或返回。",
            "H：帮助；M：音效开关；Enter：开始或继续。",
        ]),
        ("七、编译与发布", [
            r"编译打包：.\build.ps1",
            r"运行游戏：java -jar .\dist\StarMaze.jar",
            r"发布文件：dist\StarMaze.jar，Windows 下可双击 dist\run.bat。",
        ]),
        ("八、测试结果", [
            "已通过 javac 编译和 jar 打包。",
            '已通过 jar 内自检：java "-Djava.awt.headless=true" -cp dist\\StarMaze.jar com.starmaze.StarMazeSmokeTest。',
            "自检输出：StarMaze smoke test passed，并包含确定性收集星核并过关流程检查、确定性敌人碰撞失败/相位反制流程检查、确定性轨迹回溯流程检查、确定性相位结束安全恢复流程检查、EnemyController 追踪/巡逻/眩晕/无路回退检查、PathFinder A* 最短路/占用/不可达边界检查、Level 导航辅助方法检查、LevelGenerator 可达性/尺寸上限/裂隙位置检查、LevelPopulator 数量上限和安全落点检查、ScoreRules 分数和评级边界检查、SaveData 正常保存和损坏存档恢复检查、InputController 键盘绑定安装检查、BoardRenderer 静态缓存/外框/裂隙绘制检查、BackgroundRenderer 背景绘制检查、StarField 星点绘制与尺寸边界检查、PickupRenderer 拾取物绘制检查、ExitRenderer 锁定/开启出口绘制检查、PlayerRenderer 普通/相位玩家绘制检查、EnemyRenderer 方向/眩晕敌人绘制检查、MenuMouseController 鼠标悬停/点击事件检查、GameViewRenderer 渲染门面检查、MenuOverlayRenderer 菜单覆盖层分发检查、MenuChromeRenderer 遮罩和面板绘制检查、HudPanelRenderer 面板底板绘制检查、HudStatsRenderer 统计文字绘制检查、HudMeterRenderer 能量条绘制检查、HudRewindRenderer 回溯次数绘制检查、HudMessageRenderer 右侧提示绘制检查、FooterRenderer 底部提示绘制检查、TitleMenuContentRenderer 标题内容绘制检查、SimpleMenuContentRenderer 暂停/设置内容绘制检查、MenuButtonGroupRenderer 按钮组绘制检查、MenuButtonRenderer 按钮绘制检查、HelpMenuContentRenderer 帮助内容绘制检查、ResultMenuContentRenderer 结算内容绘制检查、RingEffectRenderer 环形特效绘制检查、StunEffectRenderer 眩晕点特效绘制检查、离屏 GamePanel 非空渲染检查、计时器启停检查和 GameSceneRenderer 棋盘缓存复用检查。",
        ]),
        ("九、提交文件", [
            "源码：src/main/java/com/starmaze 下的 Java 源码及 build.ps1、README.md。",
            "项目说明：docs/项目说明.doc，同时保留 docs/项目说明.docx。",
            "项目发布 Jar：dist/StarMaze.jar，另附 dist/run.bat 便于双击运行。",
        ]),
    ]
    body = []
    for title, items in sections:
        body.append(f"<h1>{escape(title)}</h1>")
        body.append("<ul>")
        for item in items:
            body.append(f"<li>{escape(item)}</li>")
        body.append("</ul>")
    html = f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>星轨迷阵 Java Swing 小游戏项目说明</title>
<style>
body {{ font-family: Calibri, 'Microsoft YaHei', sans-serif; color: #223040; margin: 48px; line-height: 1.45; }}
h1 {{ color: #2E74B5; font-size: 20px; margin-top: 22px; }}
.title {{ text-align: center; color: #0B2545; font-size: 30px; font-weight: bold; margin-bottom: 6px; }}
.subtitle {{ text-align: center; color: #566270; margin-bottom: 24px; }}
li {{ margin-bottom: 6px; }}
</style>
</head>
<body>
<div class="title">星轨迷阵 Star Maze</div>
<div class="subtitle">Java Swing 小游戏项目说明</div>
{''.join(body)}
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_html_doc()
    print(f"created {DOCX_PATH}")
    print(f"created {DOC_PATH}")
